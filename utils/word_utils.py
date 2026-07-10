# utils/word_utils.py

from docx.shared import Pt, RGBColor
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_PARAGRAPH_ALIGNMENT,
)
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ==========================================================
# FONT UTILITIES
# ==========================================================

def set_run_font(
    run,
    font_name="Times New Roman",
    font_size=11,
    bold=False,
    italic=False,
):
    """
    Apply font formatting to a Run.
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic


def set_paragraph_font(
    paragraph,
    font_name="Times New Roman",
    font_size=11,
    bold=False,
):
    """
    Apply font formatting to all runs in a paragraph.
    """
    for run in paragraph.runs:
        set_run_font(
            run,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
        )


# ==========================================================
# PARAGRAPH
# ==========================================================

def set_paragraph_spacing(
    paragraph,
    before=0,
    after=0,
    line_spacing=1.0,
):
    """
    Configure paragraph spacing.
    """
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def center_paragraph(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def left_paragraph(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def right_paragraph(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def justify_paragraph(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


# ==========================================================
# TABLE
# ==========================================================

def set_table_borders(table):
    """
    Apply borders to every cell.
    """
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def set_cell_border(cell):
    """
    Apply border around a cell.
    """

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    borders = OxmlElement("w:tcBorders")

    for edge in ("top", "left", "bottom", "right"):

        element = OxmlElement(f"w:{edge}")

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")

        borders.append(element)

    tcPr.append(borders)


# ==========================================================
# CELL
# ==========================================================

def set_cell_vertical_center(cell):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_background(cell, color="D9D9D9"):
    """
    Background color of a table cell.
    """

    tc_pr = cell._tc.get_or_add_tcPr()

    shading = OxmlElement("w:shd")

    shading.set(qn("w:fill"), color)

    tc_pr.append(shading)


def set_cell_text(
    cell,
    text,
    bold=False,
    font_size=11,
    align="left",
):
    """
    Replace cell content with formatted text.
    """

    cell.text = ""

    paragraph = cell.paragraphs[0]

    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run(text)

    set_run_font(
        run,
        font_size=font_size,
        bold=bold,
    )

    set_paragraph_spacing(paragraph)


# ==========================================================
# MULTIPLE TEXT STYLES IN SAME CELL
# ==========================================================

def add_bold_label(paragraph, label, value):
    """
    Example:

    Notes: xyz

    Objective Evidence: abc
    """

    run = paragraph.add_run(label)
    set_run_font(run, bold=True)

    run = paragraph.add_run(value)
    set_run_font(run)


def add_line(paragraph):
    paragraph.add_run("\n")


# ==========================================================
# DOCUMENT
# ==========================================================

def set_default_document_style(document):
    """
    Set default document style.
    """

    style = document.styles["Normal"]

    style.font.name = "Times New Roman"

    style._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Times New Roman",
    )

    style.font.size = Pt(11)


# ==========================================================
# TITLE
# ==========================================================

def add_document_title(
    document,
    title,
):
    """
    Add centered report title.
    """

    p = document.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(title)

    set_run_font(
        run,
        font_size=16,
        bold=True,
    )

    set_paragraph_spacing(
        p,
        before=0,
        after=10,
    )


# ==========================================================
# HEADING
# ==========================================================

def add_heading(
    document,
    text,
):
    """
    Section heading.
    """

    p = document.add_paragraph()

    run = p.add_run(text)

    set_run_font(
        run,
        font_size=12,
        bold=True,
    )

    set_paragraph_spacing(
        p,
        before=6,
        after=4,
    )


# ==========================================================
# EMPTY ROW
# ==========================================================

def add_blank_paragraph(document):
    document.add_paragraph()
from datetime import datetime

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utils.word_utils import set_run_font


# ==========================================================
# PAGE NUMBER FIELD
# ==========================================================

def add_page_number(paragraph):
    """
    Insert dynamic page number field.

    Output:
    Page No. 1
    """

    run = paragraph.add_run("Page No. ")

    set_run_font(run, bold=False)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_text)
    run._r.append(fld_end)


# ==========================================================
# HEADER
# ==========================================================

def create_header(document, company_name=""):
    """
    Creates report header.
    """

    section = document.sections[0]

    header = section.header

    paragraph = header.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(company_name)

    set_run_font(
        run,
        font_size=14,
        bold=True,
    )

    paragraph = header.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run("INTERNAL AUDIT REPORT")

    set_run_font(
        run,
        font_size=16,
        bold=True,
    )


# ==========================================================
# FOOTER
# ==========================================================

def create_footer(document):
    """
    Footer Layout

    Confidential          Page No. X          Date
    """

    section = document.sections[0]

    footer = section.footer

    table = footer.add_table(
        rows=1,
        cols=3,
        width=section.page_width
    )

    table.autofit = True

    # ---------------- Left ---------------- #

    left = table.cell(0, 0)

    p = left.paragraphs[0]

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run("Confidential")

    set_run_font(run, font_size=9)

    # ---------------- Center ---------------- #

    center = table.cell(0, 1)

    p = center.paragraphs[0]

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_number(p)

    # ---------------- Right ---------------- #

    right = table.cell(0, 2)

    p = right.paragraphs[0]

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run(
        datetime.today().strftime("%d-%m-%Y")
    )

    set_run_font(run, font_size=9)


# ==========================================================
# PAGE SETTINGS
# ==========================================================

def configure_page(document):
    """
    Configure page margins and page size.
    """

    from docx.shared import Inches

    section = document.sections[0]

    section.top_margin = Inches(0.5)

    section.bottom_margin = Inches(0.5)

    section.left_margin = Inches(0.6)

    section.right_margin = Inches(0.6)


# ==========================================================
# REPORT INFORMATION
# ==========================================================

def create_report_information(
    document,
    audit_date="",
    framework="",
    report_name=""
):
    """
    Creates

    Audit Date:

    ISO Standard Reference:

    Internal Audit Report:
    """

    table = document.add_table(
        rows=3,
        cols=2
    )

    table.style = "Table Grid"

    table.autofit = True

    labels = [
        "Audit Date",
        "ISO Standard Reference",
        "Internal Audit Report",
    ]

    values = [
        audit_date,
        framework,
        report_name,
    ]

    for i in range(3):

        cell = table.cell(i, 0)

        p = cell.paragraphs[0]

        run = p.add_run(labels[i])

        set_run_font(
            run,
            bold=True
        )

        cell = table.cell(i, 1)

        p = cell.paragraphs[0]

        run = p.add_run(values[i])

        set_run_font(run)

    document.add_paragraph()
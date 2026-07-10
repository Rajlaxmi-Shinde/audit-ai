from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================================
# MAIN GENERATOR
# ==========================================================

import os
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

from utils.word_utils import (
    set_default_document_style,
    set_cell_background,
    set_cell_text,
    add_bold_label,
    add_line,
)

from utils.footer import (
    configure_page,
    create_header,
    create_footer,
    create_report_information,
)


# ==========================================================
# CREATE MAIN FINDINGS TABLE
# ==========================================================

def create_findings_table(document):
    """
    Creates the main audit findings table.
    """

    table = document.add_table(rows=1, cols=3)

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    headers = table.rows[0].cells

    headers[0].text = "Reference of Clause or Control"

    headers[1].text = (
        "Audit Observations with Objective Evidence"
    )

    headers[2].text = "Audit Finding"

    for cell in headers:
        set_cell_background(cell, "D9D9D9")

        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in p.runs:
                run.bold = True

    return table


# ==========================================================
# ADD SINGLE FINDING
# ==========================================================

def add_finding_row(
    table,
    finding,
    notes=None,
    ):
    """
    Adds one audit finding.
    """

    row = table.add_row().cells

    # ------------------------------------

    reference = (
        f"{finding['control_id']}\n"
        f"{finding['control_name']}"
    )

    set_cell_text(row[0], reference)

    # ------------------------------------

    paragraph = row[1].paragraphs[0]

    if notes:

        add_bold_label(
            paragraph,
            "Notes:\n",
            ""
        )

        for note in notes:

            paragraph.add_run(f"• {note}\n")

        add_line(paragraph)

    add_bold_label(
        paragraph,
        "Audit Observation:\n",
        ""
    )

    paragraph.add_run(
        finding["observation"]
    )

    add_line(paragraph)

    add_line(paragraph)

    add_bold_label(
        paragraph,
        "Objective Evidence:\n",
        ""
    )

    paragraph.add_run(
        finding["objective_evidence"]
    )

    # ------------------------------------

    set_cell_text(
        row[2],
        finding["finding"],
        bold=True,
        align="center",
    )


# ==========================================================
# POPULATE FINDINGS
# ==========================================================

def populate_findings(
    table,
    findings,
    notes,
):
    """
    Insert all findings.
    """

    for index, finding in enumerate(findings):

        if index == 0:

            add_finding_row(
                table,
                finding,
                notes,
            )

        else:

            add_finding_row(
                table,
                finding,
                None,
            )


# ==========================================================
# STATIC LEGEND TABLE
# ==========================================================

def create_legend_table(document):

    document.add_paragraph()

    table = document.add_table(
        rows=4,
        cols=2,
    )

    table.style = "Table Grid"

    headers = table.rows[0].cells

    headers[0].text = "Audit Finding"

    headers[1].text = "Description"

    for cell in headers:
        set_cell_background(cell)

    values = [

        (
            "C",
            "Conformance - Positive Observation",
        ),

        (
            "NC",
            "Non-Conformance",
        ),

        (
            "S",
            "Suggestion for Improvement",
        )

    ]

    for i, item in enumerate(values):

        row = table.rows[i + 1].cells

        row[0].text = item[0]

        row[1].text = item[1]


# ==========================================================
# SIGNATURE TABLE
# ==========================================================

def create_signature_table(document):

    document.add_paragraph()

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.style = "Table Grid"

    table.cell(0, 0).text = "Auditor"

    table.cell(0, 1).text = "Auditee"

    table.cell(1, 0).text = ""

    table.cell(1, 1).text = ""

def generate_docx(report_data: dict, output_path: str):
    """
    Generate Internal Audit Report.

    Parameters
    ----------
    report_data : dict

        {
            "company_name": "",
            "audit_date": "",
            "iso_reference": "",
            "internal_audit_report": "",
            "notes": [...],
            "findings": [...]
        }

    output_path : str
    """

    logger.info("=" * 60)
    logger.info("GENERATING DOCX REPORT")
    logger.info("=" * 60)

    try:

        # --------------------------------------------------
        # Create Document
        # --------------------------------------------------

        document = Document()

        set_default_document_style(document)

        configure_page(document)

        # --------------------------------------------------
        # Header & Footer
        # --------------------------------------------------

        create_header(
            document,
            report_data.get("company_name", "")
        )

        create_footer(document)

        # --------------------------------------------------
        # Report Information
        # --------------------------------------------------

        create_report_information(
            document=document,
            audit_date=report_data.get(
                "audit_date",
                datetime.today().strftime("%d-%m-%Y")
            ),
            framework=report_data.get(
                "iso_reference",
                ""
            ),
            report_name=report_data.get(
                "internal_audit_report",
                ""
            ),
        )

        # --------------------------------------------------
        # Findings Table
        # --------------------------------------------------

        findings_table = create_findings_table(document)

        findings = report_data.get(
            "findings",
            []
        )

        notes = report_data.get(
            "notes",
            []
        )

        populate_findings(
            findings_table,
            findings,
            notes,
        )

        # --------------------------------------------------
        # Legend
        # --------------------------------------------------

        create_legend_table(document)

        # --------------------------------------------------
        # Signature
        # --------------------------------------------------

        create_signature_table(document)

        # --------------------------------------------------
        # Save
        # --------------------------------------------------

        output_directory = os.path.dirname(output_path)

        if output_directory:

            os.makedirs(
                output_directory,
                exist_ok=True,
            )

        document.save(output_path)

        logger.info(
            "Report saved successfully."
        )

        logger.info(
            "Location : %s",
            output_path
        )

        return output_path

    except Exception as e:

        logger.exception(
            "DOCX generation failed."
        )

        raise RuntimeError(
            "Unable to generate audit report."
        ) from e